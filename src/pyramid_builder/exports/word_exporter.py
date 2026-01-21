"""
Word (DOCX) export functionality for strategic pyramids.

Generates professional Word documents with formatting, tables, and structure.
"""

from typing import Optional
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from ..models.pyramid import StrategyPyramid


class WordExporter:
    """Export pyramids to Word (DOCX) format with professional formatting."""

    def __init__(self, pyramid: StrategyPyramid):
        """
        Initialize exporter.

        Args:
            pyramid: StrategyPyramid to export
        """
        self.pyramid = pyramid
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Set up custom styles for the document."""
        styles = self.doc.styles

        # Heading styles are built-in, but we can customize colors
        # Add a custom style for tier headings
        try:
            tier_style = styles.add_style('TierHeading', WD_STYLE_TYPE.PARAGRAPH)
            tier_font = tier_style.font
            tier_font.size = Pt(14)
            tier_font.bold = True
            tier_font.color.rgb = RGBColor(31, 119, 180)  # Blue color
        except:
            # Style might already exist
            pass

    def _add_vision_statements(self, heading_level=2):
        """Add vision statements to the document (handles new multi-statement structure)."""
        if not self.pyramid.vision or not self.pyramid.vision.statements:
            return

        self.doc.add_heading('Our Purpose', level=heading_level)

        for stmt in self.pyramid.vision.get_statements_ordered():
            # Add statement type as bold label
            p = self.doc.add_paragraph()
            run = p.add_run(f"{stmt.statement_type.value.title()}")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(31, 119, 180)

            # Add statement content on new line
            p2 = self.doc.add_paragraph()
            run_content = p2.add_run(stmt.statement)
            run_content.italic = True
            run_content.font.size = Pt(11)
            p2.paragraph_format.left_indent = Inches(0.25)
            p2.space_after = Pt(12)

    def export(
        self,
        filepath: str,
        audience: str = "leadership",
        include_cover_page: bool = True,
    ) -> Path:
        """
        Export pyramid to Word file.

        Args:
            filepath: Where to save the Word file
            audience: Target audience (executive, leadership, detailed, team)
            include_cover_page: Include a cover page

        Returns:
            Path to created file
        """
        if include_cover_page:
            self._add_cover_page()

        if audience == "executive":
            self._generate_executive_summary()
        elif audience == "team":
            self._generate_team_cascade()
        elif audience == "detailed":
            self._generate_detailed_strategy()
        else:  # leadership (default)
            self._generate_leadership_document()

        filepath_obj = Path(filepath)
        self.doc.save(str(filepath_obj))
        return filepath_obj

    def _add_cover_page(self):
        """Add a professional cover page."""
        # Title
        title = self.doc.add_heading(self.pyramid.metadata.project_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = self.doc.add_paragraph(self.pyramid.metadata.organization)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(16)
        subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

        # Spacer
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Metadata
        meta_table = self.doc.add_table(rows=4, cols=2)
        meta_table.style = 'Light Grid Accent 1'

        meta_table.rows[0].cells[0].text = "Created by"
        meta_table.rows[0].cells[1].text = self.pyramid.metadata.created_by

        meta_table.rows[1].cells[0].text = "Version"
        meta_table.rows[1].cells[1].text = self.pyramid.metadata.version

        meta_table.rows[2].cells[0].text = "Created"
        meta_table.rows[2].cells[1].text = self.pyramid.metadata.created_at.strftime('%d %B %Y')

        meta_table.rows[3].cells[0].text = "Last Modified"
        meta_table.rows[3].cells[1].text = self.pyramid.metadata.last_modified.strftime('%d %B %Y at %H:%M')

        # Page break
        self.doc.add_page_break()

    def _generate_executive_summary(self):
        """Generate 1-page executive summary."""
        self.doc.add_heading('Executive Summary', level=1)

        # Our Purpose
        # Vision/Mission/Belief statements
        self._add_vision_statements(heading_level=2)

        # Strategic Focus
        if self.pyramid.strategic_drivers:
            self.doc.add_heading('Strategic Focus', level=2)
            for driver in self.pyramid.strategic_drivers:
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run(f"{driver.name}: ").bold = True
                p.add_run(driver.description)

        # Key Commitments
        if self.pyramid.iconic_commitments:
            self.doc.add_heading('Key Commitments', level=2)
            # Show top 5 by horizon
            for horizon in ["H1", "H2", "H3"]:
                commitments = [c for c in self.pyramid.iconic_commitments if c.horizon.value == horizon][:2]
                if commitments:
                    for commitment in commitments:
                        p = self.doc.add_paragraph(style='List Bullet')
                        p.add_run(commitment.name).bold = True
                        if commitment.target_date:
                            p.add_run(f" ({commitment.target_date})")

    def _generate_leadership_document(self):
        """Generate full leadership document (3-5 pages)."""
        # Table of Contents
        self.doc.add_heading('Strategic Pyramid', level=1)
        self.doc.add_paragraph(f"Organisation: {self.pyramid.metadata.organization}")
        self.doc.add_paragraph(f"Last updated: {self.pyramid.metadata.last_modified.strftime('%d %B %Y')}")
        self.doc.add_paragraph()

        # SECTION 1: PURPOSE
        self.doc.add_heading('Section 1: Purpose', level=1)
        self.doc.add_paragraph('Why we exist and what matters to us').italic = True

        # Vision/Mission/Belief statements
        self._add_vision_statements(heading_level=2)

        if self.pyramid.values:
            self.doc.add_heading('Our Values', level=2)
            for value in self.pyramid.values:
                p = self.doc.add_paragraph()
                run = p.add_run(value.name)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(31, 119, 180)

                if value.description:
                    p2 = self.doc.add_paragraph()
                    p2.add_run(value.description)
                    p2.paragraph_format.left_indent = Inches(0.25)
                    p2.space_after = Pt(8)

        # SECTION 2: STRATEGY
        self.doc.add_page_break()
        self.doc.add_heading('Section 2: Strategy', level=1)
        self.doc.add_paragraph('How we will succeed').italic = True

        if self.pyramid.behaviours:
            self.doc.add_heading('Our Behaviours', level=2)
            for behaviour in self.pyramid.behaviours:
                self.doc.add_paragraph(behaviour.statement, style='List Bullet')

        if self.pyramid.strategic_drivers:
            self.doc.add_heading('Strategic Drivers', level=2)
            for driver in self.pyramid.strategic_drivers:
                self.doc.add_heading(driver.name, level=3)
                self.doc.add_paragraph(driver.description)

                # Show intents for this driver
                intents = [i for i in self.pyramid.strategic_intents if i.driver_id == driver.id]
                if intents:
                    self.doc.add_paragraph('What success looks like:').bold = True
                    for intent in intents:
                        p = self.doc.add_paragraph(style='List Bullet')
                        p.add_run(intent.statement).italic = True

        if self.pyramid.enablers:
            self.doc.add_heading('Enablers', level=2)
            p_intro = self.doc.add_paragraph('What makes our strategy possible')
            p_intro.runs[0].italic = True
            p_intro.runs[0].font.color.rgb = RGBColor(100, 100, 100)
            p_intro.space_after = Pt(12)

            for enabler in self.pyramid.enablers:
                p = self.doc.add_paragraph()
                run = p.add_run(enabler.name)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(31, 119, 180)

                if enabler.enabler_type:
                    run_type = p.add_run(f"  ")
                    run_type = p.add_run(enabler.enabler_type)
                    run_type.italic = True
                    run_type.font.size = Pt(10)
                    run_type.font.color.rgb = RGBColor(100, 100, 100)

                p2 = self.doc.add_paragraph()
                p2.add_run(enabler.description)
                p2.paragraph_format.left_indent = Inches(0.25)
                p2.space_after = Pt(8)

        # SECTION 3: EXECUTION
        self.doc.add_page_break()
        self.doc.add_heading('Section 3: Execution', level=1)
        self.doc.add_paragraph('Our iconic commitments').italic = True

        if self.pyramid.iconic_commitments:
            # Group by horizon
            for horizon in ["H1", "H2", "H3"]:
                commitments = [c for c in self.pyramid.iconic_commitments if c.horizon.value == horizon]
                if commitments:
                    horizon_name = {
                        "H1": "H1 (0-12 months)",
                        "H2": "H2 (12-24 months)",
                        "H3": "H3 (24-36 months)"
                    }[horizon]

                    self.doc.add_heading(horizon_name, level=2)

                    for commitment in commitments:
                        # Get primary driver name
                        driver = self.pyramid.get_driver_by_id(commitment.primary_driver_id)
                        driver_name = driver.name if driver else "Not specified"

                        self.doc.add_heading(commitment.name, level=3)

                        # Add details table with cleaner formatting
                        num_rows = 1
                        if commitment.target_date:
                            num_rows += 1
                        if commitment.owner:
                            num_rows += 1

                        details_table = self.doc.add_table(rows=num_rows, cols=2)
                        details_table.style = 'Light Grid Accent 1'

                        row_idx = 0
                        details_table.rows[row_idx].cells[0].text = "Primary Driver"
                        details_table.rows[row_idx].cells[1].text = driver_name

                        if commitment.target_date:
                            row_idx += 1
                            details_table.rows[row_idx].cells[0].text = "Target Date"
                            details_table.rows[row_idx].cells[1].text = commitment.target_date

                        if commitment.owner:
                            row_idx += 1
                            details_table.rows[row_idx].cells[0].text = "Owner"
                            details_table.rows[row_idx].cells[1].text = commitment.owner

                        # Make first column bold
                        for row in details_table.rows:
                            row.cells[0].paragraphs[0].runs[0].font.bold = True

                        self.doc.add_paragraph()
                        self.doc.add_paragraph(commitment.description)

                        # Show secondary alignments
                        if commitment.secondary_alignments:
                            secondary_drivers = []
                            for alignment in commitment.secondary_alignments:
                                sec_driver = self.pyramid.get_driver_by_id(alignment.target_id)
                                if sec_driver:
                                    secondary_drivers.append(sec_driver.name)
                            if secondary_drivers:
                                p = self.doc.add_paragraph()
                                p.add_run("Also contributes to: ").italic = True
                                p.add_run(", ".join(secondary_drivers)).italic = True

                        self.doc.add_paragraph()  # Spacing

        # Distribution Analysis
        if self.pyramid.iconic_commitments:
            self.doc.add_page_break()
            self.doc.add_heading('Distribution Analysis', level=1)

            distribution = self.pyramid.get_distribution_by_driver()
            total = sum(distribution.values())

            # Create distribution table
            dist_table = self.doc.add_table(rows=len(distribution) + 1, cols=3)
            dist_table.style = 'Medium Grid 1 Accent 1'

            # Header row
            dist_table.rows[0].cells[0].text = "Strategic Driver"
            dist_table.rows[0].cells[1].text = "Commitments"
            dist_table.rows[0].cells[2].text = "% of Total"

            # Data rows
            for idx, (driver_name, count) in enumerate(distribution.items(), 1):
                percentage = (count / total * 100) if total > 0 else 0
                dist_table.rows[idx].cells[0].text = driver_name
                dist_table.rows[idx].cells[1].text = str(count)
                dist_table.rows[idx].cells[2].text = f"{percentage:.0f}%"

    def _generate_detailed_strategy(self):
        """Generate detailed strategy pack (10-15 pages) with all relationships."""
        # Start with leadership document
        self._generate_leadership_document()

        # Add team objectives if present
        if self.pyramid.team_objectives:
            self.doc.add_page_break()
            self.doc.add_heading('Team Objectives', level=1)

            # Group by team
            teams = {}
            for obj in self.pyramid.team_objectives:
                if obj.team_name not in teams:
                    teams[obj.team_name] = []
                teams[obj.team_name].append(obj)

            for team_name, objectives in teams.items():
                self.doc.add_heading(team_name, level=2)

                for obj in objectives:
                    self.doc.add_heading(obj.name, level=3)
                    self.doc.add_paragraph(obj.description)

                    if obj.metrics:
                        self.doc.add_paragraph('Success Metrics:').bold = True
                        for metric in obj.metrics:
                            self.doc.add_paragraph(metric, style='List Bullet')

                    # Show relationships (NEW: supports commitment OR intent)
                    relationships = []
                    if obj.primary_commitment_id:
                        commitment = self.pyramid.get_commitment_by_id(obj.primary_commitment_id)
                        if commitment:
                            relationships.append(f"Commitment: {commitment.name}")

                    if obj.primary_intent_id:
                        intent = self.pyramid.get_intent_by_id(obj.primary_intent_id)
                        if intent:
                            relationships.append(f"Intent: {intent.statement[:50]}...")

                    if relationships:
                        p = self.doc.add_paragraph()
                        p.add_run(f"Supports: {' | '.join(relationships)}").italic = True

                    self.doc.add_paragraph()  # Spacing

        # Add individual objectives if present (NEW in v0.4.0)
        if self.pyramid.individual_objectives:
            self.doc.add_page_break()
            self.doc.add_heading('Individual Objectives', level=1)

            # Group by individual
            individuals = {}
            for obj in self.pyramid.individual_objectives:
                if obj.individual_name not in individuals:
                    individuals[obj.individual_name] = []
                individuals[obj.individual_name].append(obj)

            for individual_name, objectives in individuals.items():
                self.doc.add_heading(individual_name, level=2)

                for obj in objectives:
                    self.doc.add_heading(obj.name, level=3)
                    self.doc.add_paragraph(obj.description)

                    if obj.success_criteria:
                        self.doc.add_paragraph('Success Criteria:').bold = True
                        for criterion in obj.success_criteria:
                            self.doc.add_paragraph(criterion, style='List Bullet')

                    # Show which team objectives this supports (NEW relationship)
                    if obj.team_objective_ids:
                        team_objs = []
                        for team_id in obj.team_objective_ids:
                            team_obj = next((to for to in self.pyramid.team_objectives if to.id == team_id), None)
                            if team_obj:
                                team_objs.append(f"{team_obj.team_name}: {team_obj.name}")

                        if team_objs:
                            p = self.doc.add_paragraph()
                            p.add_run('Supports Team Objectives: ').bold = True
                            p.add_run(' | '.join(team_objs)).italic = True

                    self.doc.add_paragraph()  # Spacing

    def _generate_team_cascade(self):
        """Generate team cascade view showing line of sight."""
        self.doc.add_heading('Team Cascade View', level=1)
        self.doc.add_paragraph('Line of sight from purpose to team objectives').italic = True
        self.doc.add_paragraph()

        # Vision
        # Vision/Mission/Belief statements
        self._add_vision_statements(heading_level=2)

        # For each driver, show the cascade
        for driver in self.pyramid.strategic_drivers:
            self.doc.add_page_break()
            self.doc.add_heading(driver.name, level=1)
            self.doc.add_paragraph(driver.description)

            # Intents
            intents = [i for i in self.pyramid.strategic_intents if i.driver_id == driver.id]
            if intents:
                self.doc.add_heading('What Success Looks Like', level=2)
                for intent in intents:
                    p = self.doc.add_paragraph(style='List Bullet')
                    p.add_run(intent.statement).italic = True

            # Commitments
            commitments = self.pyramid.get_commitments_by_driver(driver.id, primary_only=True)
            if commitments:
                self.doc.add_heading('Our Commitments', level=2)
                for commitment in commitments:
                    p = self.doc.add_paragraph(style='List Bullet 2')
                    p.add_run(commitment.name).bold = True
                    if commitment.target_date:
                        p.add_run(f" ({commitment.target_date})")

                    # Show related team objectives
                    related_objectives = [
                        obj for obj in self.pyramid.team_objectives
                        if obj.primary_commitment_id == commitment.id
                    ]
                    if related_objectives:
                        for obj in related_objectives:
                            sub_p = self.doc.add_paragraph(style='List Bullet 3')
                            sub_p.add_run(f"{obj.team_name}: {obj.name}")
